from datetime import timedelta
import requests
import requests_cache
import json
from bs4 import BeautifulSoup
from bs4.element import NavigableString
from jinja2 import Environment, PackageLoader, select_autoescape
from docxtpl import DocxTemplate
from mailmerge import MailMerge
from docx import Document
from docx.text.paragraph import Paragraph


class Planning:
    # Aim: grab data from page such as: "https://planning-schemes.app.planning.vic.gov.au/Victoria%20Planning%20Provisions/ordinance/3870230"
    index_url = "https://api.vicplanning.app/planning/v2/schemes/vpp"
    ordinance_url = "https://api.vicplanning.app/planning/v2/schemes/vpp/ordinances/"

    def __init__(self, clause_name, sub_clause_name):
        self.session = requests_cache.CachedSession("schemes", expire_after=timedelta(days=7))
        self.index_json = self.getUrl(self.index_url)
        self.clause_name = clause_name
        self.sub_clause_name = sub_clause_name

    def getUrl(self, url):
        return json.loads(self.session.get(url).text)

    def getOrdinanceID(self):
        for clause in self.index_json["clauses"]:
            if clause["title"] == self.clause_name:
                for sub_clause in clause["subClauses"]:
                    if sub_clause["title"] == self.sub_clause_name:
                        self.ordinance_id = sub_clause["ordinanceID"]
                        return self.ordinance_id

    def getOrdinanceJson(self):
        if not hasattr(self, "ordinance_id"):
            _ = self.getOrdinanceID()

        self.clause_json = self.getUrl(self.ordinance_url + self.ordinance_id)
        return self.clause_json

    def printOrdinance(self):
        if not hasattr(self, "clause_json"):
            _ = self.getOrdinanceJson()

        print("<html><body>")
        print(f"<h1>{self.clause_name} - {self.sub_clause_name}</h1>")
        print(self.clause_json["content"])
        for section in self.clause_json["ordinanceSections"]:
            print(f"<h2>{section['title']}</h2>")
            print(section["content"])
            print("")
        print("</body></html>")

    def parseOrdinance(self):
        if not hasattr(self, "clause_json"):
            _ = self.getOrdinanceJson()
        sections = {section["title"]: section["content"] for section in self.clause_json["ordinanceSections"]}

        self.ordinance_sections = {}
        for title, section in sections.items():
            soup = BeautifulSoup(section, "html.parser")
            self.ordinance_sections[title] = self._parse_ord_section(soup)

    def renderOrdinance(self):
        if not hasattr(self, "ordinance_sections"):
            _ = self.parseOrdinance()

        def table(tbl):
            if "caption" in tbl.keys():
                caption = f"### {tbl['caption']} \n \n"
            else:
                caption = ""
            header = "|"
            for item in tbl["header"]:
                header += f" {match_type(item)}  |"
            header += "\n |"
            for _ in tbl["header"]:
                header += " --- |"

            body = ""
            for row in tbl["body"]:
                body += "|"
                for cell in row:
                    body += f" {match_type(cell)} |"
                body += "\n"

            return f"\n{caption} {header} \n {body}\n"

        def ul(ul, multi_line=False, indent=0):
            line_list = ""
            multi_list = "\n"
            for li in ul:
                line_list += match_type(li, multi_line) + " <br />"
                multi_list += (" " * indent) + f" - {match_type(li, multi_line, indent=(indent+4))} \n"
            line_list += ""
            multi_list += "\n"
            return (line_list, multi_list)

        def match_type(item, multi_line=False, indent=0):
            if type(item) is list:
                items = []
                for sub_item in item:
                    items.append(match_type(sub_item, multi_line, indent=indent))
                if multi_line:
                    return "\n  ".join(items)
                else:
                    return "<br />".join(items)

            if "table" in item.keys():
                render = table(item["table"])
            elif "ul" in item.keys():
                if multi_line:
                    (_, render) = ul(item["ul"], multi_line, indent=indent)
                else:
                    (render, _) = ul(item["ul"], multi_line)
            elif "p" in item.keys():
                render = item["p"]
            else:
                print("Unexpected element")
                render = ""

            return render

        renders = {}
        for name, ordinance in self.ordinance_sections.items():
            items = []
            for item in ordinance:
                contents = []
                for obj in item["content"]:
                    content = match_type(obj, multi_line=True)
                    contents.append(content)
                if "title" in item.keys():
                    items.append((item["title"], contents))
                else:
                    items.append((_, contents))
            renders[name] = items

        return renders

    def _parse_children(self, elem):
        cell_children = []
        for cell_child in elem.children:
            sub_child = self._parse_elem(cell_child)
            if sub_child:
                cell_children.append(sub_child)
        if len(cell_children) == 1:
            return cell_children[0]
        elif len(cell_children) > 1:
            return cell_children

    def _parse_elem(self, elem):
        if elem.name == "ul":
            points = []
            for point in elem.children:
                if point.name == "li":
                    sub_points = self._parse_children(point)
                    if sub_points:
                        points.append({"li": sub_points})
            return {"ul": points}
        elif elem.name == "table":
            table = {"header": [], "body": []}
            for col in elem.tbody.tr.find_all("th"):
                sub_points = self._parse_children(col)
                if sub_points:
                    table["header"].append(sub_points)
            if elem.find("caption"):
                table["caption"] = elem.caption.get_text()

            for row in elem.tbody.tr.next_siblings:
                if row.name == "tr":
                    row_contents = []
                    for col in row.find_all("td"):
                        if col.name == "td":
                            sub_points = self._parse_children(col)
                            if sub_points:
                                row_contents.append(sub_points)
                    table["body"].append(row_contents)
            return {"table": table}
        elif elem.name == "p":
            text = elem.get_text()
            if not text == "":
                return {"p": text}
        elif elem.name == "br":
            return None
        elif not type(elem) == NavigableString:
            print(f"Unknown tag: {elem}")
            return elem.text
        else:
            return None

    def _parse_ord_section(self, soup):
        rules = []
        current_rule = {"content": []}
        for child in soup.children:
            if child.name == "h3":
                if len(current_rule["content"]) > 0:
                    rules.append(current_rule)
                    current_rule = {"content": []}
                current_rule["title"] = child.get_text()
            else:
                content = self._parse_elem(child)
                if content:
                    current_rule["content"].append(content)

        if len(current_rule["content"]) > 0:
            rules.append(current_rule)

        if rules:
            return rules
        else:
            print("Failed to find rule")


def docx_tpl(content):
    doc = DocxTemplate("App/templates/template.docx")
    doc.render(content)
    doc.save("output.docx")


def docx(content):
    doc = Document()

    for name, division in content.items():
        doc.add_heading(name, level=2)
        for section in division:
            if "title" in section.keys():
                doc.add_heading(section["title"], level=3)
            for content in section["content"]:
                doc = parse_elem(doc, content)

    doc.save("test.docx")


def get_text(obj):
    if type(obj) is list:
        items = []
        for item in obj:
            items.append(get_text(item))
        return " ".join(items)
    else:
        items = []
        for item in obj:
            if item == "p":
                items.append(obj[item])
            elif item == "ul" or item == "li":
                items.append(get_text(obj[item]))
            elif item == "table":
                print("Table inside list")
        return " ".join(items)


def cycle_elem(doc, obj, indent=0, paragraph=None):
    if type(obj) is list:
        cycle_elem(doc, obj[0], indent, paragraph)
        for item in obj[1:]:
            cycle_elem(doc, item, indent=indent)
    else:
        parse_elem(doc, obj, indent, paragraph)


def parse_elem(doc, obj, indent=0, paragraph=None):
    if "p" in obj.keys():
        if paragraph is not None:
            paragraph.add_run(obj["p"])
        else:
            doc.add_paragraph(obj["p"])
    elif "ul" in obj.keys():
        for elem in obj["ul"]:
            doc = parse_elem(doc, elem, indent)
    elif "li" in obj.keys():
        styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
        if indent >= len(styles):
            print(f"Indent level {indent} not supported")
            style = styles[-1]
        else:
            style = styles[indent]

        paragraph = doc.add_paragraph("", style=style)
        cycle_elem(doc, obj["li"], indent=indent + 1, paragraph=paragraph)
    elif "table" in obj.keys():
        if "caption" in obj["table"].keys():
            doc.add_heading(obj["table"]["caption"], level=4)
        row_num, col_num = len(obj["table"]["body"]) + 1, len(obj["table"]["header"])
        table = doc.add_table(rows=row_num, cols=col_num)
        for row in range(row_num):
            for cell in range(col_num):
                if row == 0:
                    table.rows[row].cells[cell].text = get_text(obj["table"]["header"][cell])
                else:
                    if len(obj["table"]["body"]) > row - 1 and len(obj["table"]["body"][row - 1]) > cell:
                        paragraph = table.rows[row].cells[cell].paragraphs[0]
                        cycle_elem(table.rows[row].cells[cell], obj["table"]["body"][row - 1][cell], indent, paragraph)
        table.style = "LightShading-Accent1"

        # print(obj["table"])

    return doc


def md(content):
    env = Environment(loader=PackageLoader("App"), autoescape=select_autoescape())
    template = env.get_template("template.md")
    render = template.render(content)
    with open("output.md", "w") as f:
        f.write(render)


if __name__ == "__main__":
    planning = Planning("32 RESIDENTIAL ZONES", "32.08 GENERAL RESIDENTIAL ZONE")

    # with open("dump2.json", "w") as f:
    #    json.dump(planning.renderOrdinance(), f)

    planning.parseOrdinance()
    # print(json.dumps(planning.ordinance_sections))
    docx(planning.ordinance_sections)

    # content = {"subdivisions": planning.renderOrdinance()}

    # with MailMerge("planning_report.docx") as doc:
    #    print(doc.get_merge_fields())

    # md(content)
